import tkinter as tk
from tkinter import ttk, filedialog
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import openpyxl 
import openpyxl.drawing
import openpyxl.drawing.image
from openpyxl import Workbook
from openpyxl.drawing.image import Image
from openpyxl.utils import get_column_letter
import openpyxl.styles
from openpyxl.styles import Alignment, Border, Side, Font, PatternFill, PatternFill
import openpyxl
from io import BytesIO
from controllers.tooltip_controller import ToolTipController
import locale

# ตั้งค่าภาษาและภูมิภาค
locale.setlocale(locale.LC_ALL, 'th_TH.UTF-8')

class ConnewView(ttk.Frame):

    def __init__(self, controller, app):
        super().__init__(app)
        self.controller = controller
        self.items = []

        self.style = ttk.Style()
        self.style.configure("My.TFrame", background='#C0E4F6')

        # Window
        self.main_frame =  ttk.Frame(self)

        self.label_profile = ttk.Label(self, borderwidth=1, relief="ridge", text = "", font=('Tohama', 10, 'bold'))
        self.label_profile.grid(row=0, column=0, columnspan=2, sticky='NSWE')
        self.label_profile.configure(anchor='center', background='#C0E4F6')

        # สร้าง Frame เพิ่มเติมเพื่อควบคุมการวางและกรอบสามเหลี่ยม
        frame_labels = ttk.Frame(self, borderwidth=1, relief="ridge", style="My.TFrame")
        frame_labels.grid(row=3, column=1, sticky='NSWE')

        # Label แรก
        self.label_totalcf = ttk.Label(frame_labels, justify='center', text="ค่าคาร์บอนฟุตพริ้นท์รวม", font=('Tohama', 10, 'bold'))
        self.label_totalcf.grid(row=0, column=0, padx=10, pady=20, sticky='E')
        self.label_totalcf.configure(anchor='e', background='#C0E4F6')

        # Label ที่สอง
        self.label_cf = ttk.Label(frame_labels, justify='center', text="", width=15, font=('bold'))
        self.label_cf.grid(row=0, column=1, padx=15, pady=20)
        self.label_cf.configure(anchor='center', background='#C0E4F6')

        # Label ที่สาม
        self.label_unit = ttk.Label(frame_labels, justify='center', text="KgCO2eq", font=('bold'))
        self.label_unit.grid(row=0, column=2, padx=5, pady=20)
        self.label_unit.configure(anchor='w', background='#C0E4F6')
        
        self.return_button = ttk.Button(self, text="Return to Profile", command=self.back)
        self.return_button.grid(row=4, rowspan=2, column=1, padx=10, pady=10, ipadx=10, ipady=10, sticky = 'W')

        self.export_button = ttk.Button(self, text="Export to Excel", command=self.export)
        self.export_button.grid(row=4, rowspan=2, column=1, padx=10, pady=10, ipadx=10, ipady=10, sticky = 'E')

        # Frame Input Treeview
        frame_input = ttk.Frame(self, borderwidth=1, relief="ridge", style="My.TFrame")
        frame_input.grid(row=2, column=0, sticky='NSWE')

        # TREE VIEW
        self.input_treeview = ttk.Treeview(frame_input, columns=("Name", "Carbon"), show="headings") # , "Unit"
        self.input_treeview.heading("Name", text="ชื่อ")
        self.input_treeview.column("Name", width=370, stretch=True)
        text_width = len("ค่าคาร์บอนเทียบเท่า (Y)")  # คำนวณความยาวของข้อความ
        self.input_treeview.heading("Carbon", text="ค่าคาร์บอน (Y)")
        self.input_treeview.column("Carbon", width=text_width * 5)
        # self.input_treeview.heading("Unit", text="หน่วย")
        # self.input_treeview.column("Unit", width=60, stretch=True)
        self.input_treeview.grid(row=0, column=0)

        # สร้าง Scrollbar แนวแกน Y
        scroll_y = ttk.Scrollbar(frame_input, orient='vertical', command=self.input_treeview.yview)
        self.input_treeview.configure(yscrollcommand=scroll_y.set)
        scroll_y.grid(row=0, column=1, sticky='NS')

        self.grid_rowconfigure(0, weight=1)

         # Frame Process Treeview
        frame_process = ttk.Frame(self, borderwidth=1, relief="ridge", style="My.TFrame")
        frame_process.grid(row=2, column=1, sticky='NSWE')

        self.process_treeview = ttk.Treeview(frame_process, columns=("Name", "Carbon"), show="headings")
        self.process_treeview.heading("Name", text="ชื่อ (X)")
        self.process_treeview.column("Name", width=370, stretch=True)
        text_width = len("ค่าคาร์บอนเทียบเท่า (Y)")  # คำนวณความยาวของข้อความ
        self.process_treeview.heading("Carbon", text="ค่าคาร์บอน (Y)")
        self.process_treeview.column("Carbon", width=text_width * 5)
        # self.process_treeview.heading("Unit", text="หน่วย")
        # self.process_treeview.column("Unit", width=60, stretch=True)
        self.process_treeview.grid(row=0, column=0)

        # สร้าง Scrollbar แนวแกน Y
        scroll_y = ttk.Scrollbar(frame_process, orient='vertical', command=self.process_treeview.yview)
        self.process_treeview.configure(yscrollcommand=scroll_y.set)
        scroll_y.grid(row=0, column=1, sticky='NS')

        self.grid_rowconfigure(0, weight=1)

        # Frame Breakeven_point
        frame_break = ttk.Frame(self, borderwidth=1, relief="ridge", style='My.TFrame')
        frame_break.grid(row=3, rowspan=3, column=0, sticky='NSWE')
        
        self.totalcost = ttk.Label(frame_break, text = "ต้นทุนรวม", font=('Tohama', 10, 'bold'))
        self.totalcost.grid(row=0, column=0, padx=30, pady=5, sticky='W')
        self.totalcost.configure(anchor='w', background='#C0E4F6')

        self.revenue = ttk.Label(frame_break, text = "รายได้", font=('Tohama', 10, 'bold'))
        self.revenue.grid(row=1, column=0, padx=30, pady=5, sticky='W')
        self.revenue.configure(anchor='w', background='#C0E4F6')
        
        self.profit = ttk.Label(frame_break, text = "กำไร", font=('Tohama', 10, 'bold'))
        self.profit.grid(row=2, column=0, padx=30, pady=5, sticky='W')
        self.profit.configure(anchor='w', background='#C0E4F6')
        
        self.breakeven = ttk.Label(frame_break, text = "ปริมาณผลิตที่จุดคุ้มทุน", font=('Tohama', 10, 'bold'))
        self.breakeven.grid(row=3, column=0, padx=30, pady=5, sticky='W')
        self.breakeven.configure(anchor='w', background='#C0E4F6')

        self.efficiency = ttk.Label(frame_break, text = "ประสิทธิภาพการผลิต", font=('Tohama', 10, 'bold'))
        self.efficiency.grid(row=4, column=0, padx=30, pady=5, sticky='W')
        self.efficiency.configure(anchor='w', background='#C0E4F6')

        self.add_totalcost = ttk.Label(frame_break, text = "", font=('Tohama', 10, 'bold'))
        self.add_totalcost.grid(row=0, column=1, padx=40, pady=5, sticky='E')
        self.add_totalcost.configure(anchor='w', background='#C0E4F6')

        self.add_revenue = ttk.Label(frame_break, text = "", font=('Tohama', 10, 'bold'))
        self.add_revenue.grid(row=1, column=1, padx=40, pady=5, sticky='E')
        self.add_revenue.configure(anchor='w', background='#C0E4F6')

        self.add_profit = ttk.Label(frame_break, text = "", font=('Tohama', 10, 'bold'))
        self.add_profit.grid(row=2, column=1, padx=40, pady=5, sticky='E')
        self.add_profit.configure(anchor='w', background='#C0E4F6')

        self.add_breakeven = ttk.Label(frame_break, text = "", font=('Tohama', 10, 'bold'))
        self.add_breakeven.grid(row=3, column=1, padx=40, pady=5, sticky='E')
        self.add_breakeven.configure(anchor='w', background='#C0E4F6')

        self.add_efficiency = ttk.Label(frame_break, text = "", font=('Tohama', 10, 'bold'))
        self.add_efficiency.grid(row=4, column=1, padx=40, pady=5, sticky='E')
        self.add_efficiency.configure(anchor='w', background='#C0E4F6')

        self.unit_totalcost = ttk.Label(frame_break, text = "บาท", font=('Tohama', 10, 'bold'))
        self.unit_totalcost.grid(row=0, column=2, padx=40, pady=5)
        self.unit_totalcost.configure(anchor='e', background='#C0E4F6')

        self.unit_revenue = ttk.Label(frame_break, text = "บาท", font=('Tohama', 10, 'bold'))
        self.unit_revenue.grid(row=1, column=2, padx=40, pady=5)
        self.unit_revenue.configure(anchor='e', background='#C0E4F6')

        self.unit_profit = ttk.Label(frame_break, text = "บาท", font=('Tohama', 10, 'bold'))
        self.unit_profit.grid(row=2, column=2, padx=40, pady=5)
        self.unit_profit.configure(anchor='e', background='#C0E4F6')

        self.unit_breakeven = ttk.Label(frame_break, text = "หน่วย", font=('Tohama', 10, 'bold'))
        self.unit_breakeven.grid(row=3, column=2, padx=40, pady=5)
        self.unit_breakeven.configure(anchor='e', background='#C0E4F6')

        self.unit_efficiency = ttk.Label(frame_break, text = "%", font=('Tohama', 10, 'bold'))
        self.unit_efficiency.grid(row=4, column=2, padx=40, pady=5)
        self.unit_efficiency.configure(anchor='e', background='#C0E4F6')
        
        # self.output_treeview = ttk.Treeview(self, columns=("Name", "Carbon", "Unit"), show="headings")
        # self.output_treeview.heading("Name", text="ชื่อ")
        # self.output_treeview.column("Name", width=310)
        # self.output_treeview.heading("Carbon", text="ค่าคาร์บอนเทียบเท่า")
        # self.output_treeview.column("Carbon", width=95)
        # self.output_treeview.heading("Unit", text="หน่วย")
        # self.output_treeview.column("Unit", width=60)
        # self.output_treeview.grid(row=3, rowspan=2, column=2, padx=5, pady=5)

        self.add_button_tooltips()
    
    def add_button_tooltips(self):
        ToolTipController(self.export_button, "ส่งออกไปยัง Excel")
        ToolTipController(self.return_button, "กลับไปยังหน้าหลัก")

    def setInputGraph(self, items):
        # สร้างรูปภาพ Matplotlib
        figure = Figure(figsize=(6.5, 3.25), dpi=70)
        subplot = figure.add_subplot(111)

        # ข้อมูลเส้นกราฟ
        x = []
        y = []
        for item in items:
            x.append(item[0])
            y.append(float(item[1]))

        # ตั้งค่าฟอนต์
        subplot.tick_params(axis='x', labelrotation=0, labelsize=10, colors='white')

        # สร้างกราฟด้วยสีส้ม
        subplot.plot(x, y, color='#AFD7F6')

        # เพิ่มจุดบ่งบอกตำแหน่งด้วยสีส้ม
        subplot.scatter(x, y, color='#AFD7F6')

        # เพิ่มตัวเลขกำกับ
        for i, txt in enumerate(y):
            subplot.annotate(f'{txt}', (x[i], y[i]), textcoords="offset points", xytext=(0,10), ha='center', fontsize=8)

        # เพิ่มชื่อกราฟ
        subplot.set_title('Carbon footprint of Input', fontsize=12, fontweight='bold')

        # เพิ่มป้ายกำกับแกน
        subplot.set_ylabel('KgCO2eq', fontsize=8, fontweight='bold')

        # ปรับแต่งป้ายกำกับ
        subplot.tick_params(axis='both', which='major', labelsize=8)

        # ปรับแต่งขอบ
        subplot.spines['top'].set_visible(False)
        subplot.spines['right'].set_visible(False)

        # สร้างกรอบสำหรับแคนวาสพร้อมเส้นขอบ
        frame = tk.Frame(self, highlightbackground='black', highlightthickness=1, borderwidth=1, relief="ridge")
        frame.grid(row=1, column=0, sticky='NSWE')

        # สร้างวิดเจ็ต FigureCanvasTkAgg
        canvas = FigureCanvasTkAgg(figure, master=frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    def setProcessGraph(self, items):
        # สร้างรูปภาพ Matplotlib
        figure = Figure(figsize=(6.5, 3.5), dpi=70)
        subplot = figure.add_subplot(111)

        # ข้อมูลเส้นกราฟ
        x = []
        y = []
        for item in items:
            x.append(item[0])
            y.append(float(item[1]))

        # ตั้งค่าฟอนต์
        subplot.tick_params(axis='x', labelrotation=0, labelsize=10, colors='white')

        # สร้างกราฟด้วยสีส้ม
        subplot.plot(x, y, color='#F7D07A')

        # เพิ่มจุดบ่งบอกตำแหน่งด้วยสีส้ม
        subplot.scatter(x, y, color='#F7D07A')

        # เพิ่มตัวเลขกำกับ
        for i, txt in enumerate(y):
            subplot.annotate(f'{txt}', (x[i], y[i]), textcoords="offset points", xytext=(0,10), ha='center', fontsize=8)

        # เพิ่มชื่อกราฟ
        subplot.set_title('Carbon footprint of Process', fontsize=12, fontweight='bold')

        # เพิ่มป้ายกำกับแกน
        subplot.set_ylabel('KgCO2eq', fontsize=8, fontweight='bold')

        # ปรับแต่งป้ายกำกับ
        subplot.tick_params(axis='both', which='major', labelsize=8)

        # ปรับแต่งขอบ
        subplot.spines['top'].set_visible(False)
        subplot.spines['right'].set_visible(False)

        # สร้างกรอบสำหรับแคนวาสพร้อมเส้นขอบ
        frame = tk.Frame(self, highlightbackground='black', highlightthickness=1, borderwidth=1, relief="ridge")
        frame.grid(row=1, column=1, sticky='NSWE')

        # สร้างวิดเจ็ต FigureCanvasTkAgg
        canvas = FigureCanvasTkAgg(figure, master=frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    # def setOutputGraph(self, items):
    #     # Create a Matplotlib figure
    #     figure = Figure(figsize=(6.5, 3), dpi=70)
    #     subplot = figure.add_subplot(111)
    
    #     # Query from database
    #     # Line data
    #     x = []
    #     y = []
    #     for item in items:
    #         x.append(item[0])
    #         y.append(float(item[1]))

    #     # Set font
    #     subplot.tick_params(axis='x', labelrotation=90, labelfontfamily="Tahoma")
    
    #     # Create graph
    #     subplot.plot(x, y)
    
    #     # Add grid
    #     subplot.grid(True, linestyle='--', linewidth=0.5)

    #     # Add title
    #     subplot.set_title('Output', fontsize=14, fontweight='bold')

    #     # Customize tick labels
    #     subplot.tick_params(axis='both', which='major', labelsize=8)
    
    #     # Add labels to axes
    #     subplot.set_xlabel('X Axis Label', fontsize=10)
    #     subplot.set_ylabel('Y Axis Label', fontsize=10)

    #     # # Add legend
    #     # subplot.legend(['Legend'], loc='upper right', fontsize=8)

    #     # Customize borders
    #     subplot.spines['top'].set_visible(False)
    #     subplot.spines['right'].set_visible(False)

    #     # Create a FigureCanvasTkAgg widget
    #     canvas = FigureCanvasTkAgg(figure, master=self)
    #     canvas.draw()
    #     canvas.get_tk_widget().grid(row=1, rowspan=2, column=2, padx=5, pady=5)

    def back(self):
        self.controller.back_main()
    
    def format_currency(self, value):
        return locale.currency(value, grouping=True)

    def set_breakpoint_data(self, breakpoint_data):
        if breakpoint_data:
            fixed_cost, variable_cost, number_of_units, unit_price, product_efficiency = breakpoint_data[0]
            total_cost = fixed_cost + (variable_cost * number_of_units)
            revenue = unit_price * number_of_units
            profit = revenue - total_cost
            breakeven = fixed_cost / (unit_price - variable_cost)
            
            # อัปเดตค่า total cost
            self.add_totalcost.config(text=self.format_currency(total_cost))
            # อัปเดตค่า revenue
            self.add_revenue.config(text=self.format_currency(revenue))
            # อัปเดตค่า profit
            self.add_profit.config(text=self.format_currency(profit), foreground="red" if profit < 0 else "green")
            # อัปเดตค่า breakeven
            self.add_breakeven.config(text=f"{breakeven:.2f}", foreground="red" if breakeven < 0 else "black")
            # อัปเดตค่า efficiency
            self.add_efficiency.config(text=f"{product_efficiency:.2f}")

            # จัดเก็บค่าที่คำนวณไว้ในตัวแปรชั่วคราว
            self.total_cost = total_cost
            self.revenue = revenue
            self.profit = profit
            self.breakeven = breakeven
            self.product_efficiency = product_efficiency
            
        else:
            self.add_totalcost.config(text="-", foreground="black")
            self.add_revenue.config(text="-", foreground="black")
            self.add_profit.config(text="-", foreground="black")
            self.add_breakeven.config(text="-", foreground="black")
            self.add_efficiency.config(text="-", foreground="black")

            # รีเซ็ตตัวแปรชั่วคราว
            self.total_cost = None
            self.revenue = None
            self.profit = None
            self.breakeven = None
            self.product_efficiency = None

    def setConclusion(self, profile_name, items, breakpoint_data) :
        self.input_treeview.delete(*self.input_treeview.get_children())
        self.process_treeview.delete(*self.process_treeview.get_children())
        # self.output_treeview.delete(*self.output_treeview.get_children())
        inputGraph = []
        # outputGraph = []
        processGraph = []

        total_cf = 0  # ผลรวมของค่า Carbon

        for item in items:
            category, _, name, factor, amount, _ = item
            
            # self.input_treeview
            if category == 'Material':
                processGraph.append((name, round(float(factor) * float(amount), 3)))
                self.process_treeview.insert("", "end", values=(name, round(float(factor) * float(amount), 3))) # , "KgCO2eq"
                total_cf += float(factor) * float(amount)  # เพิ่มค่า Carbon ลงในผลรวม

            elif category == 'Transpotation':
                inputGraph.append((name, round(float(factor) * float(amount), 3)))
                # outputGraph.append((name, round(float(factor) * float(amount), 3)))
                self.input_treeview.insert("", "end", values=(name, round(float(factor) * float(amount), 3)))
                # self.output_treeview.insert("", "end", values=(name, round(float(factor) * float(amount), 3), "KgCO2eq"))
                total_cf += float(factor) * float(amount)  # เพิ่มค่า Carbon ลงในผลรวม

            elif category == 'Performance':
                processGraph.append((name, round(float(factor) * float(amount), 3)))
                self.process_treeview.insert("", "end", values=(name, round(float(factor) * float(amount), 3)))
                total_cf += float(factor) * float(amount)  # เพิ่มค่า Carbon ลงในผลรวม

        self.label_cf.config(text="{:.3f}".format(total_cf))  # กำหนดผลรวมของค่า Carbon ลงใน label_cf

        self.setInputGraph(inputGraph)
        # self.setOutputGraph(outputGraph)
        self.setProcessGraph(processGraph)
        self.set_breakpoint_data(breakpoint_data)
        self.items = items
        self.label_profile.config(text=profile_name)

    def process_item_data(self, items):
        input_data, process_data = [], []
        for item in items:
            category, _, name, carbon_per, amount, unit = item
            carbon_footprint = round(float(carbon_per) * float(amount), 2)
            data = (name, carbon_per, amount, unit, carbon_footprint)
            if category == 'Material':
                process_data.append(data)
            elif category == 'Transpotation':
                input_data.append(data)
            elif category == 'Performance':
                process_data.append(data)
        return input_data, process_data

    def calculate_totals(self, data):
        return round(sum(item[4] for item in data), 2)

    def create_styles(self):
        align_center = Alignment(horizontal="center", vertical="center")
        return align_center
    
    def add_financial_summary(self, sheet, summary_row):
        align_center = self.create_styles()
        financial_data = [
            ("Total cost", self.total_cost, '฿#,##0.00', 'Baht'),
            ("Revenue", self.revenue, '฿#,##0.00', 'Baht'),
            ("Profit", self.profit, '฿#,##0.00', 'Baht'),
            ("Break-even Point", self.breakeven, '0.00', 'Unit'),
            ("Product Efficiency", self.product_efficiency, '0.00', '%')]

        for i, (label, value, num_format, suffix) in enumerate(financial_data, start=1):
            row = summary_row + i + 1
            sheet[f"A{row}"] = label
            sheet[f"A{row}"].font = Font(name='TH Sarabun New', size=12)
            cell_value = sheet[f"B{row}"]
            cell_value.value = value
            cell_value.font = Font(name='TH Sarabun New', size=12)
            cell_value.number_format = num_format
            cell_suffix = sheet[f"C{row}"]
            cell_suffix.value = suffix
            cell_suffix.font = Font(name='TH Sarabun New', size=12)
            cell_suffix.alignment = align_center
      

    def set_column_widths(self, sheet):
        for column_cells in sheet.columns:
            max_length = max(len(str(cell.value)) for cell in column_cells if cell.value)
            adjusted_width = max_length + 2
            sheet.column_dimensions[get_column_letter(column_cells[0].column)].width = adjusted_width

    def set_page_layout(self, sheet):
        sheet.page_setup.fitToWidth = 1
        sheet.page_setup.fitToHeight = 0
        sheet.page_setup.paperSize = sheet.PAPERSIZE_A4

    def create_graph(self, items, title, color):
        if not items:
            print("No data available to create the graph")
            return None
        
        figure = Figure(figsize=(6.5, 3.25), dpi=70)
        subplot = figure.add_subplot(111)
        x, y = zip(*[(item[0], float(item[4])) for item in items if len(item) >= 5])
        subplot.tick_params(axis='x', labelrotation=45, labelsize=10, colors='white')
        subplot.plot(x, y, color=color)
        subplot.scatter(x, y, color=color)
        for i, txt in enumerate(y):
            subplot.annotate(f'{txt}', (x[i], y[i]), textcoords="offset points", xytext=(0, 10), ha='center', fontsize=8)
        subplot.set_title(title, fontsize=12, fontweight='bold')
        subplot.set_ylabel('KgCO2eq', fontsize=8, fontweight='bold')
        subplot.tick_params(axis='both', which='major', labelsize=8)
        subplot.spines['top'].set_visible(False)
        subplot.spines['right'].set_visible(False)
        buffer = BytesIO()
        figure.savefig(buffer, format="png")
        buffer.seek(0)
        return buffer
    
    def insert_graphs(self, sheet, start_row, input_graph, process_graph):
        from openpyxl.drawing.image import Image

        input_graph_col = "A"
        process_graph_col = "I"  # คอลัมน์ที่ 11 นับจาก A (ห่างกัน 10 คอลัมน์)

        # แทรกกราฟของ Input
        if input_graph:
            img1 = Image(input_graph)
            sheet.add_image(img1, f"{input_graph_col}{start_row + 1}")
        
        # แทรกกราฟของ Process
        if process_graph:
            img2 = Image(process_graph)
            sheet.add_image(img2, f"{process_graph_col}{start_row + 1}")

    # def insert_graphs(self, sheet, start_row, input_graph, process_graph):
    #     input_graph_row = start_row + 2
    #     if input_graph:
    #         sheet.add_image(Image(input_graph), f"A{input_graph_row}")
    #     process_graph_row = input_graph_row + 15
    #     if process_graph:
    #         sheet.add_image(Image(process_graph), f"A{process_graph_row}")

    # def insert_graphs(self, sheet, summary_row, input_graph, process_graph):
    #     input_graph_row = summary_row + 11
    #     if input_graph:
    #         sheet.add_image(Image(input_graph), f"A{input_graph_row}")
    #     process_graph_row = input_graph_row + 15
    #     if process_graph:
    #         sheet.add_image(Image(process_graph), f"A{process_graph_row}")

    def export(self):
        if any(attr is None for attr in [self.total_cost, self.revenue, self.profit, self.breakeven, self.product_efficiency]):
            print("No data available for export")
            return

        # Get project name from label_profile
        project_name = self.label_profile.cget("text")

        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("Word files", "*.docx")],
            title="Save file as")

        if file_path:
            if file_path.endswith('.xlsx'):
                self.export_excel(file_path, project_name)
            elif file_path.endswith('.docx'):
                self.export_docx(file_path, project_name)

    def export_excel(self, file_path, project_name):
        if any(attr is None for attr in [self.total_cost, self.revenue, self.profit, self.breakeven, self.product_efficiency]):
            print("No data available for export")
            return

        input_data, process_data = self.process_item_data(self.items)
        total_input_carbon_footprint = self.calculate_totals(input_data)
        total_process_carbon_footprint = self.calculate_totals(process_data)
        total_carbon_footprint = round(total_input_carbon_footprint + total_process_carbon_footprint, 2)

        align_center = self.create_styles()

        wb = Workbook()
        sheet = wb.active
        sheet.title = "Carbon Footprint Report"

        # Create report header
        sheet.merge_cells("A1:E1")
        sheet["A1"].value = "Carbon Footprint Report"
        sheet["A1"].alignment = align_center
        sheet["A1"].font = Font(name='TH Sarabun New', size=14, bold=True)

        # Create project name header
        sheet.merge_cells("A2:E2")
        sheet["A2"].value = f"Project Name: {project_name}"
        # sheet["A2"].alignment = align_center
        sheet["A2"].font = Font(name='TH Sarabun New', size=12, bold=True)

        # Create table headers
        headers = ["Name", "Emission Factor", "Amount", "Unit", "Carbon Footprint"]
        for col_num, header in enumerate(headers, start=1):
            cell = sheet.cell(row=3, column=col_num)
            cell.value = header
            cell.alignment = align_center
            cell.font = Font(name='TH Sarabun New', size=12, bold=True)

        # Fill in data rows
        row_num = 4
        for data in input_data + process_data:
            for col_num, value in enumerate(data, start=1):
                cell = sheet.cell(row=row_num, column=col_num)
                cell.value = value
                cell.font = Font(name='TH Sarabun New', size=12)
                if col_num in [2, 3, 5]:
                    cell.number_format = '0.00'
            row_num += 1

        # Add summary data
        summary_row = row_num + 1
        sheet[f"A{summary_row}"] = "Total Carbon Footprint"
        sheet[f"A{summary_row}"].font = Font(name='TH Sarabun New', size=12)
        sheet[f"B{summary_row}"] = total_carbon_footprint
        sheet[f"B{summary_row}"].font = Font(name='TH Sarabun New', size=12)
        sheet[f"B{summary_row}"].number_format = '0.00'
        sheet[f"C{summary_row}"] = "KgCO2eq"
        sheet[f"C{summary_row}"].alignment = align_center
        sheet[f"C{summary_row}"].font = Font(name='TH Sarabun New', size=12)

        self.add_financial_summary(sheet, summary_row)
        self.set_column_widths(sheet)
        self.set_page_layout(sheet)

        # Create new sheet for graphs
        graph_sheet = wb.create_sheet(title="Graphs")
        graph_sheet.merge_cells("A1:E1")
        graph_sheet["A1"].value = "Carbon Footprint Graphs"
        graph_sheet["A1"].alignment = align_center
        graph_sheet["A1"].font = Font(name='TH Sarabun New', size=14, bold=True)

        input_graph = self.create_graph(input_data, 'Carbon footprint of Input', '#AFD7F6')
        process_graph = self.create_graph(process_data, 'Carbon footprint of Process', '#F7D07A')
        self.insert_graphs(graph_sheet, 2, input_graph, process_graph)

        input_row_num = 20
        input_headers = ["Name", "Carbon Footprint"]
        for col_num, header in enumerate(input_headers, start=1):
            cell = graph_sheet.cell(row=input_row_num, column=col_num)
            cell.value = header
            cell.alignment = align_center
            cell.font = Font(name='TH Sarabun New', size=12, bold=True)

        input_row_num += 1
        for data in input_data:
            name, carbon_footprint = data[0], data[4]
            cell_name = graph_sheet.cell(row=input_row_num, column=1)
            cell_name.value = name
            cell_name.font = Font(name='TH Sarabun New', size=12)
            cell_footprint = graph_sheet.cell(row=input_row_num, column=2)
            cell_footprint.value = carbon_footprint
            cell_footprint.font = Font(name='TH Sarabun New', size=12)
            cell_footprint.number_format = '0.00'
            input_row_num += 1

        for col in range(1, 3):
            max_length = 0
            column = graph_sheet.column_dimensions[chr(64 + col)]
            for row in range(20, input_row_num):
                cell = graph_sheet.cell(row=row, column=col)
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            adjusted_width = (max_length + 2)
            column.width = adjusted_width

        process_row_num = input_row_num + 2
        process_headers = ["Name", "Carbon Footprint"]
        for col_num, header in enumerate(process_headers, start=1):
            cell = graph_sheet.cell(row=process_row_num, column=col_num + 10)
            cell.value = header
            cell.alignment = align_center
            cell.font = Font(name='TH Sarabun New', size=12, bold=True)

        process_row_num += 1
        for data in process_data:
            name, carbon_footprint = data[0], data[4]
            cell_name = graph_sheet.cell(row=process_row_num, column=11)
            cell_name.value = name
            cell_name.font = Font(name='TH Sarabun New', size=12)
            cell_footprint = graph_sheet.cell(row=process_row_num, column=12)
            cell_footprint.value = carbon_footprint
            cell_footprint.font = Font(name='TH Sarabun New', size=12)
            cell_footprint.number_format = '0.00'
            process_row_num += 1

        for col in range(11, 13):
            max_length = 0
            column = graph_sheet.column_dimensions[chr(64 + col)]
            for row in range(input_row_num + 2, process_row_num):
                cell = graph_sheet.cell(row=row, column=col)
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            adjusted_width = (max_length + 2)
            column.width = adjusted_width

        if file_path:
            wb.save(file_path)
            print(f"File saved at: {file_path}")
        else:
            print("File save canceled")

    def export_docx(self, file_path, project_name):
        document = Document()

        # Set landscape orientation
        section = document.sections[-1]
        section.orientation = WD_ORIENTATION.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Add report heading
        heading = document.add_heading('Carbon Footprint Report', level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in heading.runs:
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True

        # Add project name heading
        subheading = document.add_heading(f"Project Name: {project_name}", level=1)
        for run in subheading.runs:
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True

        input_data, process_data = self.process_item_data(self.items)
        total_input_carbon_footprint = self.calculate_totals(input_data)
        total_process_carbon_footprint = self.calculate_totals(process_data)
        total_carbon_footprint = round(total_input_carbon_footprint + total_process_carbon_footprint, 2)

        # # Add input data heading
        # p_input_heading = document.add_paragraph('Input')
        # p_input_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # run_input_heading = p_input_heading.runs[0]
        # run_input_heading.font.name = 'TH Sarabun New'
        # run_input_heading.font.size = Pt(18)
        # run_input_heading.bold = True

        # # Add table for input data
        # table_input = document.add_table(rows=1, cols=5)
        # hdr_cells_input = table_input.rows[0].cells
        # headers = ['Name', 'Emission Factor', 'Amount', 'Unit', 'Carbon Footprint']
        # for cell, header in zip(hdr_cells_input, headers):
        #     cell.text = header
        #     for paragraph in cell.paragraphs:
        #         for run in paragraph.runs:
        #             run.font.name = 'TH Sarabun New'
        #             run.font.size = Pt(18)
        #             run.bold = True

        # for item in input_data:
        #     row_cells = table_input.add_row().cells
        #     for cell, value in zip(row_cells, item):
        #         cell.text = str(value)
        #         for paragraph in cell.paragraphs:
        #             for run in paragraph.runs:
        #                 run.font.name = 'TH Sarabun New'
        #                 run.font.size = Pt(16)

        # # Add process data heading
        # p_process_heading = document.add_paragraph('Process')
        # p_process_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # run_process_heading = p_process_heading.runs[0]
        # run_process_heading.font.name = 'TH Sarabun New'
        # run_process_heading.font.size = Pt(18)
        # run_process_heading.bold = True

        # # Add table for process data
        # table_process = document.add_table(rows=1, cols=5)
        # hdr_cells_process = table_process.rows[0].cells
        # for cell, header in zip(hdr_cells_process, headers):
        #     cell.text = header
        #     for paragraph in cell.paragraphs:
        #         for run in paragraph.runs:
        #             run.font.name = 'TH Sarabun New'
        #             run.font.size = Pt(18)
        #             run.bold = True

        # for item in process_data:
        #     row_cells = table_process.add_row().cells
        #     for cell, value in zip(row_cells, item):
        #         cell.text = str(value)
        #         for paragraph in cell.paragraphs:
        #             for run in paragraph.runs:
        #                 run.font.name = 'TH Sarabun New'
        #                 run.font.size = Pt(16)

        # Add table for data
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        headers = ['Name', 'Emission Factor', 'Amount', 'Unit', 'Carbon Footprint']
        for cell, header in zip(hdr_cells, headers):
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'TH Sarabun New'
                    run.font.size = Pt(18)
                    run.bold = True

        for item in input_data + process_data:
            row_cells = table.add_row().cells
            for cell, value in zip(row_cells, item):
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'TH Sarabun New'
                        run.font.size = Pt(16)

        # Add summary data in a single paragraph
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        run_left = p.add_run('Total Carbon Footprint:')
        run_left.font.name = 'TH Sarabun New'
        run_left.font.size = Pt(16)
        run_left.bold = True
        run_left.font.color.rgb = RGBColor(0, 0, 0)
        
        tab = p.add_run('\t\t\t\t')
        tab.font.name = 'TH Sarabun New'
        tab.font.size = Pt(16)
        
        run_center = p.add_run(f'{total_carbon_footprint}')
        run_center.font.name = 'TH Sarabun New'
        run_center.font.size = Pt(16)
        run_center.font.color.rgb = RGBColor(0, 0, 0)
        
        tab = p.add_run('\t\t\t\t\t\t')
        tab.font.name = 'TH Sarabun New'
        tab.font.size = Pt(16)
        
        run_right = p.add_run('KgCO2eq')
        run_right.font.name = 'TH Sarabun New'
        run_right.font.size = Pt(16)
        run_right.font.color.rgb = RGBColor(0, 0, 0)

        # Add "Financial Summary" heading on a new page
        document.add_page_break()
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run('Financial Summary')
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(18)
        run.bold = True

        # Add financial summary in a table
        financial_data = [
            ("Total cost", f"{self.total_cost:,.2f}", 'Baht'),
            ("Revenue", f"{self.revenue:,.2f}", 'Baht'),
            ("Profit", f"{self.profit:,.2f}", 'Baht'),
            ("Break-even Point", f"{self.breakeven:.2f}", 'Units'),
            ("Product Efficiency", f"{self.product_efficiency:.2f}", '%')]
        
        table = document.add_table(rows=0, cols=3)
        for label, value, suffix in financial_data:
            row_cells = table.add_row().cells
            label_cell = row_cells[0].paragraphs[0].add_run(label)
            label_cell.font.name = 'TH Sarabun New'
            label_cell.font.size = Pt(16)
            label_cell.bold = True

            value_paragraph = row_cells[1].paragraphs[0]
            value_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            value_cell = value_paragraph.add_run(value)
            value_cell.font.name = 'TH Sarabun New'
            value_cell.font.size = Pt(16)

            suffix_paragraph = row_cells[2].paragraphs[0]
            suffix_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            suffix_cell = suffix_paragraph.add_run(suffix)
            suffix_cell.font.name = 'TH Sarabun New'
            suffix_cell.font.size = Pt(16)

        # Add graphs on a new page
        document.add_page_break()
        input_graph = self.create_graph(input_data, 'Carbon footprint of Input', '#AFD7F6')
        process_graph = self.create_graph(process_data, 'Carbon footprint of Process', '#F7D07A')
        # # Add graphs
        # input_graph = self.create_graph(input_data, 'Carbon footprint of Input', '#AFD7F6')
        # process_graph = self.create_graph(process_data, 'Carbon footprint of Process', '#F7D07A')

        # Add graph images to document and center them
        if input_graph:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            run.add_picture(input_graph, width=Inches(6))
            last_paragraph = document.paragraphs[-1]
            for run in last_paragraph.runs:
                run.font.name = 'TH Sarabun New'

        if process_graph:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            run.add_picture(process_graph, width=Inches(6))
            last_paragraph = document.paragraphs[-1]
            for run in last_paragraph.runs:
                run.font.name = 'TH Sarabun New'

        if file_path:
            document.save(file_path)
            print(f"File saved at: {file_path}")
        else:
            print("File save canceled")