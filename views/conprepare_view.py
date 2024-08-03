import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import openpyxl 
import openpyxl.drawing
import openpyxl.drawing.image
from openpyxl import Workbook
from openpyxl.drawing.image import Image
from openpyxl.utils import get_column_letter
import openpyxl.styles
from openpyxl.styles import Alignment, Border, Side, Font, PatternFill, PatternFill
import openpyxl
from io import BytesIO
from controllers.tooltip_controller import ToolTipController
import locale

# ตั้งค่าภาษาและภูมิภาค
locale.setlocale(locale.LC_ALL, 'th_TH.UTF-8')

class ConprepareView(ttk.Frame):

    def __init__(self, controller, app):
        super().__init__(app)
        self.controller = controller
        self.items = []

        self.style = ttk.Style()
        self.style.configure("My.TFrame", background='#C0E4F6')

        # Window
        self.main_frame =  ttk.Frame(self)

        self.label_profile1 = ttk.Label(self, borderwidth=1, relief="ridge"
                                        , text = "", justify='center', font=('Tohama', 12, 'bold'), background='#C0E4F6') 
        self.label_profile1.grid(row=1, column=0, sticky='NSWE')
        self.label_profile1.configure(anchor='center')

        self.label_profile2 = ttk.Label(self, borderwidth=1, relief="ridge"
                                        , text = "", justify='center', font=('Tohama', 12, 'bold'), background='#C0E4F6')
        self.label_profile2.grid(row=1, column=1, sticky='NSWE')
        self.label_profile2.configure(anchor='center')

        # Frame CF
        frame_cf = ttk.Frame(self, borderwidth=1, relief="ridge")
        frame_cf.grid(row=5, column=2, sticky='NSWE')

        self.label_percentcf = ttk.Label(frame_cf, borderwidth=1, relief="ridge", justify='center', text = "   ส่วนต่างคาร์บอนฟุตพริ้นท์", font=('Tohama', 11, 'bold'))
        self.label_percentcf.grid(row=0, column=0, sticky='NSWE')
        self.label_percentcf.configure(anchor='w', background='#C0E4F6')

        self.label_cf = ttk.Label(frame_cf, borderwidth=1, relief="ridge", justify='center', text = "", font=('Tohama', 11, 'bold'))
        self.label_cf.grid(row=0, column=1, sticky = 'NSWE')
        self.label_cf.configure(anchor='center', background='white')

        self.dif_revenue = ttk.Label(frame_cf, borderwidth=1, relief="ridge", justify='center', text = "   ส่วนต่างต้นทุนรวม", font=('Tohama', 11, 'bold'))
        self.dif_revenue.grid(row=1, column=0, sticky='NSWE')
        self.dif_revenue.configure(anchor='w', background='#C0E4F6')

        self.add_dif_revenue = ttk.Label(frame_cf, borderwidth=1, relief="ridge", justify='center', text = "", font=('Tohama', 11, 'bold'))
        self.add_dif_revenue.grid(row=1, column=1, sticky='NSWE')
        self.add_dif_revenue.configure(anchor='center', background='white')

        self.dif_percent = ttk.Label(frame_cf, borderwidth=1, relief="ridge", justify='center', text = "   ส่วนต่างกำไร", font=('Tohama', 11, 'bold'))
        self.dif_percent.grid(row=2, column=0, sticky='NSWE')
        self.dif_percent.configure(anchor='w', background='#C0E4F6')

        self.add_dif_percent = ttk.Label(frame_cf, borderwidth=1, relief="ridge", justify='center', text = "", font=('Tohama', 11, 'bold'))
        self.add_dif_percent.grid(row=2, column=1, sticky = 'NSWE')
        self.add_dif_percent.configure(anchor='center', background='white')

        # Button
        self.return_button = ttk.Button(self, text="Back to Main", command=self.back)
        self.return_button.grid(row=5, column=0, padx=10, pady=10, ipadx=10, ipady=10, sticky = 'W')

        self.export_button = ttk.Button(self, text="Export", command=self.export)
        self.export_button.grid(row=5, column=0, padx=10, pady=10, ipadx=10, ipady=10, sticky = 'E')

        # Frame Profile1
        frame_profile1 = ttk.Frame(self, borderwidth=1, relief="ridge", style="My.TFrame")
        frame_profile1.grid(row=2, rowspan=3, column=0, sticky='NSWE')

        self.profile1_treeview = ttk.Treeview(frame_profile1, columns=("Name", "Carbon"), show="headings")
        self.profile1_treeview.heading("Name", text="Name (X)")
        self.profile1_treeview.column("Name", width=370, stretch=True)
        text_width = len("คาร์บอนเทียบเท่า (Y)")  # คำนวณความยาวของข้อความ
        self.profile1_treeview.heading("Carbon", text="Carbon (Y)")
        self.profile1_treeview.column("Carbon", width=text_width * 5)
        # self.profile1_treeview.heading("Unit", text="หน่วย")
        # self.profile1_treeview.column("Unit", width=60, stretch=True)
        self.profile1_treeview.grid(row=0, column=0)

        # สร้าง Scrollbar แนวแกน Y
        scroll_y = ttk.Scrollbar(frame_profile1, orient='vertical', command=self.profile1_treeview.yview)
        self.profile1_treeview.configure(yscrollcommand=scroll_y.set)
        scroll_y.grid(row=0, column=1, sticky='NS')

        self.grid_rowconfigure(0, weight=1)

        # Frame Profile2
        frame_profile2 = ttk.Frame(self, borderwidth=1, relief="ridge", style="My.TFrame")
        frame_profile2.grid(row=2, rowspan=3, column=1, sticky='NSWE')

        self.profile2_treeview = ttk.Treeview(frame_profile2, columns=("Name", "Carbon"), show="headings") # , "Unit"
        self.profile2_treeview.heading("Name", text="Name (X)")
        self.profile2_treeview.column("Name", width=370, stretch=True)
        text_width = len("คาร์บอนเทียบเท่า (Y)")  # คำนวณความยาวของข้อความ
        self.profile2_treeview.heading("Carbon", text="Carbon (Y)")
        self.profile2_treeview.column("Carbon", width=text_width * 5)
        # self.profile2_treeview.heading("Unit", text="หน่วย")
        # self.profile2_treeview.column("Unit", width=60, stretch=True)
        self.profile2_treeview.grid(row=0, column=0)

        # สร้าง Scrollbar แนวแกน Y
        scroll_y = ttk.Scrollbar(frame_profile2, orient='vertical', command=self.profile2_treeview.yview)
        self.profile2_treeview.configure(yscrollcommand=scroll_y.set)
        scroll_y.grid(row=0, column=1, sticky='NS')

        self.grid_rowconfigure(0, weight=1)

        # Frame Breakeven_point
        frame_break = ttk.Frame(self, borderwidth=1, relief="ridge")
        frame_break.grid(row=0, column=2, sticky='NS')

        self.break_profile1 = ttk.Label(frame_break, text = "", borderwidth=1, relief="ridge", font=('Tohama', 12, 'bold'))
        self.break_profile1.grid(row=0, column=0, columnspan=3, sticky='NSWE')
        self.break_profile1.configure(anchor='center', background='#C0E4F6')

        self.totalcost = ttk.Label(frame_break, text = "Total cost", font=('Tohama', 11, 'bold'))
        self.totalcost.grid(row=1, column=0, padx=10, pady=10, sticky='W')
        self.totalcost.configure(anchor='w', background='white')

        # self.revenue = ttk.Label(frame_break, text = "รายได้")
        # self.revenue.grid(row=2, column=0, padx=10, pady=11, sticky='W')
        # self.revenue.configure(anchor='w', background='white')
        
        self.profit = ttk.Label(frame_break, text = "Profit", font=('Tohama', 11, 'bold'))
        self.profit.grid(row=2, column=0, padx=10, pady=10, sticky='W')
        self.profit.configure(anchor='w', background='white')
        
        self.breakeven = ttk.Label(frame_break, text = "Break-even point", font=('Tohama', 11, 'bold'))
        self.breakeven.grid(row=3, column=0, padx=10, pady=10, sticky='W')
        self.breakeven.configure(anchor='w', background='white')

        # self.efficiency = ttk.Label(frame_break, text = "ประสิทธิภาพการผลิต")
        # self.efficiency.grid(row=4, column=0, padx=10, pady=11, sticky='W')
        # self.efficiency.configure(anchor='w', background='white')

        self.add_totalcost = ttk.Label(frame_break, text = "", font=('Tohama', 11, 'bold'))
        self.add_totalcost.grid(row=1, column=1, padx=10, pady=10, sticky='E')
        self.add_totalcost.configure(anchor='w', background='white')

        # self.add_revenue = ttk.Label(frame_break, text = "")
        # self.add_revenue.grid(row=2, column=1, padx=10, pady=11, sticky='E')
        # self.add_revenue.configure(anchor='w', background='#FFD10A')

        self.add_profit = ttk.Label(frame_break, text = "", font=('Tohama', 11, 'bold'))
        self.add_profit.grid(row=2, column=1, padx=10, pady=10, sticky='E')
        self.add_profit.configure(anchor='w', background='white')

        self.add_breakeven = ttk.Label(frame_break, text = "", font=('Tohama', 11, 'bold'))
        self.add_breakeven.grid(row=3, column=1, padx=10, pady=10, sticky='E')
        self.add_breakeven.configure(anchor='w', background='white')

        # self.add_efficiency = ttk.Label(frame_break, text = "")
        # self.add_efficiency.grid(row=4, column=1, padx=10, pady=11, sticky='E')
        # self.add_efficiency.configure(anchor='w', background='#FFD10A')

        self.unit_totalcost = ttk.Label(frame_break, text = "Baht", font=('Tohama', 11, 'bold'))
        self.unit_totalcost.grid(row=1, column=2, padx=10, pady=10)
        self.unit_totalcost.configure(anchor='e', background='white')

        # self.unit_revenue = ttk.Label(frame_break, text = "บาท")
        # self.unit_revenue.grid(row=2, column=2, padx=10, pady=11)
        # self.unit_revenue.configure(anchor='e', background='white')

        self.unit_profit = ttk.Label(frame_break, text = "Baht", font=('Tohama', 11, 'bold'))
        self.unit_profit.grid(row=2, column=2, padx=10, pady=10)
        self.unit_profit.configure(anchor='e', background='white')

        self.unit_breakeven = ttk.Label(frame_break, text = "Unit", font=('Tohama', 11, 'bold'))
        self.unit_breakeven.grid(row=3, column=2, padx=10, pady=10)
        self.unit_breakeven.configure(anchor='e', background='white')

        # self.unit_efficiency = ttk.Label(frame_break, text = "%")
        # self.unit_efficiency.grid(row=4, column=2, padx=10, pady=11)
        # self.unit_efficiency.configure(anchor='e', background='white')

        self.break_profile2 = ttk.Label(frame_break, text = "", borderwidth=1, relief="ridge", font=('Tohama', 12, 'bold'))
        self.break_profile2.grid(row=4, column=0, columnspan=3, sticky='NSWE')
        self.break_profile2.configure(anchor='center', background='#C0E4F6')
        
        self.totalcost2 = ttk.Label(frame_break, text = "Total cost", font=('Tohama', 11, 'bold'))
        self.totalcost2.grid(row=5, column=0, padx=10, pady=10, sticky='W')
        self.totalcost2.configure(anchor='w', background='white')

        # self.revenue2 = ttk.Label(frame_break, text = "รายได้")
        # self.revenue2.grid(row=2, column=0, padx=10, pady=11, sticky='W')
        # self.revenue2.configure(anchor='w', background='white')
        
        self.profit2 = ttk.Label(frame_break, text = "Profit", font=('Tohama', 11, 'bold'))
        self.profit2.grid(row=6, column=0, padx=10, pady=10, sticky='W')
        self.profit2.configure(anchor='w', background='white')
        
        self.breakeven2 = ttk.Label(frame_break, text = "Break-even point", font=('Tohama', 11, 'bold'))
        self.breakeven2.grid(row=7, column=0, padx=10, pady=10, sticky='W')
        self.breakeven2.configure(anchor='w', background='white')

        # self.efficiency2 = ttk.Label(frame_break, text = "ประสิทธิภาพการผลิต")
        # self.efficiency2.grid(row=4, column=0, padx=10, pady=11, sticky='W')
        # self.efficiency2.configure(anchor='w', background='white')

        self.add_totalcost2 = ttk.Label(frame_break, text = "", font=('Tohama', 11, 'bold'))
        self.add_totalcost2.grid(row=5, column=1, padx=10, pady=10, sticky='E')
        self.add_totalcost2.configure(anchor='w', background='white')

        # self.add_revenue2 = ttk.Label(frame_break, text = "")
        # self.add_revenue2.grid(row=2, column=1, padx=10, pady=11, sticky='E')
        # self.add_revenue2.configure(anchor='w', background='#FFD10A')

        self.add_profit2 = ttk.Label(frame_break, text = "", font=('Tohama', 11, 'bold'))
        self.add_profit2.grid(row=6, column=1, padx=10, pady=10, sticky='E')
        self.add_profit2.configure(anchor='w', background='white')

        self.add_breakeven2 = ttk.Label(frame_break, text = "", font=('Tohama', 11, 'bold'))
        self.add_breakeven2.grid(row=7, column=1, padx=10, pady=10, sticky='E')
        self.add_breakeven2.configure(anchor='w', background='white')

        # self.add_efficiency2 = ttk.Label(frame_break, text = "")
        # self.add_efficiency2.grid(row=4, column=1, padx=10, pady=11, sticky='E')
        # self.add_efficiency2.configure(anchor='w', background='#FFD10A')

        self.unit_totalcost2 = ttk.Label(frame_break, text = "Baht", font=('Tohama', 11, 'bold'))
        self.unit_totalcost2.grid(row=5, column=2, padx=10, pady=10)
        self.unit_totalcost2.configure(anchor='e', background='white')

        # self.unit_revenue2 = ttk.Label(frame_break, text = "บาท")
        # self.unit_revenue2.grid(row=2, column=2, padx=10, pady=11)
        # self.unit_revenue2.configure(anchor='e', background='white')

        self.unit_profit2 = ttk.Label(frame_break, text = "Baht", font=('Tohama', 11, 'bold'))
        self.unit_profit2.grid(row=6, column=2, padx=10, pady=10)
        self.unit_profit2.configure(anchor='e', background='white')

        self.unit_breakeven2 = ttk.Label(frame_break, text = "Unit", font=('Tohama', 11, 'bold'))
        self.unit_breakeven2.grid(row=7, column=2, padx=10, pady=10)
        self.unit_breakeven2.configure(anchor='e', background='white')

        # self.unit_efficiency2 = ttk.Label(frame_break, text = "%")
        # self.unit_efficiency2.grid(row=4, column=2, padx=10, pady=11)
        # self.unit_efficiency2.configure(anchor='e', background='white')

        # Frame Recommend (รวมทั้งสาม frame เป็นกรอบเดียว)
        frame_recommend = ttk.Frame(self, borderwidth=1, relief="ridge")
        frame_recommend.grid(row=1, rowspan=4, column=2, sticky='NSWE')

        self.label_carbon_footprint = ttk.Label(frame_recommend, borderwidth=1, relief="ridge", font=('Tohama', 11, 'bold')
                                                , text="พิจารณาที่คาร์บอนฟุตพริ้นท์", anchor='center', background='#C0E4F6')
        self.label_carbon_footprint.grid(row=0, column=0, sticky='NSWE')

        self.label_compare_summary = ttk.Label(frame_recommend, justify='center', font=('Tohama', 11, 'bold')
                                               , text="", anchor='center', background='white')
        self.label_compare_summary.grid(row=1, column=0, padx=10, pady=10, sticky='NSWE')

        self.label_total_cost = ttk.Label(frame_recommend, borderwidth=1, relief="ridge", font=('Tohama', 11, 'bold')
                                          , text="พิจารณาต้นทุน", anchor='center', background='#C0E4F6')
        self.label_total_cost.grid(row=2, column=0, sticky='NSWE')

        self.label_total_cost_summary = ttk.Label(frame_recommend, justify='center', font=('Tohama', 11, 'bold')
                                                  , text="", anchor='center', background='white')
        self.label_total_cost_summary.grid(row=3, column=0, padx=10, pady=10, sticky='NSWE')

        self.label_profit = ttk.Label(frame_recommend, borderwidth=1, relief="ridge", font=('Tohama', 11, 'bold')
                                      , text="พิจารณากำไร", anchor='center', background='#C0E4F6')
        self.label_profit.grid(row=4, column=0, sticky='NSWE')

        self.label_profit_summary = ttk.Label(frame_recommend, justify='center', font=('Tohama', 11, 'bold')
                                              , text="", anchor='center', background='white')
        self.label_profit_summary.grid(row=5, column=0, padx=10, pady=10, sticky='NSWE')

        # Configure the grid to make it flexible
        self.grid_rowconfigure(2, weight=1)
        self.grid_rowconfigure(3, weight=1)
        self.grid_rowconfigure(4, weight=1)
        self.grid_columnconfigure(2, weight=1)

        frame_cf.grid_rowconfigure(0, weight=1)
        frame_cf.grid_rowconfigure(1, weight=1)
        frame_cf.grid_rowconfigure(2, weight=1)
        frame_cf.grid_columnconfigure(0, weight=1)
        frame_cf.grid_columnconfigure(1, weight=1)

        frame_profile1.grid_rowconfigure(0, weight=1)
        frame_profile1.grid_columnconfigure(0, weight=1)
        frame_profile1.grid_columnconfigure(1, weight=0)

        frame_profile2.grid_rowconfigure(0, weight=1)
        frame_profile2.grid_columnconfigure(0, weight=1)
        frame_profile2.grid_columnconfigure(1, weight=0)

        frame_recommend.grid_rowconfigure(0, weight=1)
        frame_recommend.grid_rowconfigure(1, weight=1)
        frame_recommend.grid_rowconfigure(2, weight=1)
        frame_recommend.grid_rowconfigure(3, weight=1)
        frame_recommend.grid_rowconfigure(4, weight=1)
        frame_recommend.grid_rowconfigure(5, weight=1)
        frame_recommend.grid_columnconfigure(0, weight=1)

        self.add_button_tooltips()

        self.items = []

        self.rawmats_1 = []
        self.transpots_1 = []
        self.performances_1 = []
        self.rawmats_2 = []
        self.transpots_2 = []
        self.performances_2 = []

    def add_button_tooltips(self):
        ToolTipController(self.export_button, "ส่งออกไปยัง Excel")
        ToolTipController(self.return_button, "กลับไปยังหน้าหลัก")
        
    def back(self):
        self.controller.back_main()

    def format_currency(self, value):
        return locale.currency(value, grouping=True)
    
    def set_breakpoint_data(self, profile1, profile2, breakpoint_data1, breakpoint_data2):
        revenue1 = None
        revenue2 = None
        profit1 = 0.0
        profit2 = 0.0

        if breakpoint_data1:
            fixed_cost, variable_cost, number_of_units, unit_price, product_efficiency = breakpoint_data1[0]
            total_cost1 = fixed_cost + (variable_cost * number_of_units)
            revenue1 = unit_price * number_of_units
            profit1 = revenue1 - total_cost1
            breakeven1 = fixed_cost / (unit_price - variable_cost)
            
            self.add_totalcost.config(text=self.format_currency(total_cost1))
            # self.add_revenue.config(text=self.format_currency(revenue1))
            self.add_profit.config(text=self.format_currency(profit1), foreground="red" if profit1 < 0 else "green")
            self.add_breakeven.config(text=f"{breakeven1:.2f}",foreground="red" if breakeven1 < 0 else "black")
            # self.add_efficiency.config(text=f"{product_efficiency:.2f}")

            self.total_cost1 = total_cost1
            self.revenue1 = revenue1
            self.profit1 = profit1
            self.breakeven1 = breakeven1
            self.product_efficiency1 = product_efficiency

        else:
            self.add_totalcost.config(text="-", foreground="black")
            # self.add_revenue.config(text="-", foreground="black")
            self.add_profit.config(text="-", foreground="black")
            self.add_breakeven.config(text="-", foreground="black")
            # self.add_efficiency.config(text="-", foreground="black")

            self.total_cost1 = None
            self.revenue1 = None
            self.profit1 = None
            self.breakeven1 = None
            self.product_efficiency1 = None

        if breakpoint_data2:
            fixed_cost, variable_cost, number_of_units, unit_price, product_efficiency = breakpoint_data2[0]
            total_cost2 = fixed_cost + (variable_cost * number_of_units)
            revenue2 = unit_price * number_of_units
            profit2 = revenue2 - total_cost2
            breakeven2 = fixed_cost / (unit_price - variable_cost)

            self.add_totalcost2.config(text=self.format_currency(total_cost2))
            # self.add_revenue2.config(text=self.format_currency(revenue2))
            self.add_profit2.config(text=self.format_currency(profit2), foreground="red" if profit2 < 0 else "green")
            self.add_breakeven2.config(text=f"{breakeven2:.2f}", foreground="red" if breakeven2 < 0 else "black")
            # self.add_efficiency2.config(text=f"{product_efficiency:.2f}")

            self.total_cost2 = total_cost2
            self.revenue2 = revenue2
            self.profit2 = profit2
            self.breakeven2 = breakeven2
            self.product_efficiency2 = product_efficiency

        else:
            self.add_totalcost2.config(text="-", foreground="black")
            # self.add_revenue2.config(text="-", foreground="black")
            self.add_profit2.config(text="-", foreground="black")
            self.add_breakeven2.config(text="-", foreground="black")
            # self.add_efficiency2.config(text="-", foreground="black")

            self.total_cost2 = None
            self.revenue2 = None
            self.profit2 = None
            self.breakeven2 = None
            self.product_efficiency2 = None

        # คำนวณและแสดงความแตกต่างเปอร์เซ็นต์ในกำไรหากมีทั้งสองกำไรที่พร้อมใช้งาน
        if profit1 is not None and profit2 is not None:
            percentage_difference = ((profit2 - profit1) / profit1) * 100
            if percentage_difference >= 0:
                self.add_dif_percent.config(
                    text=f"+{percentage_difference:.2f}%",
                    foreground="green")
            else:
                self.add_dif_percent.config(
                    text=f"{percentage_difference:.2f}%",
                    foreground="red")

            # แสดงผลต่าง % ของ revenue1 กับ revenue2 ใน add_dif_revenue
            revenue_difference = ((revenue2 - revenue1) / revenue1) * 100
            if revenue_difference >= 0:
                self.add_dif_revenue.config(
                    text=f"+{revenue_difference:.2f}%",
                    foreground="green")
            else:
                self.add_dif_revenue.config(
                    text=f"{revenue_difference:.2f}%",
                    foreground="red")
        else:
            self.add_dif_percent.config(text="-", foreground="black")
            self.add_dif_revenue.config(text="-", foreground="black")

        # เรียกฟังก์ชัน setBreakevenGraph เพื่อสร้างกราฟ โดยส่ง profile1, profile2, profit1 และ profit2 ไปให้
        self.setBreakevenGraph(profile1, profile2, total_cost1, total_cost2, profit1, profit2)

    def setCompareGraph(self, profile1, profile2):
        # สร้างภาพ Matplotlib
        figure = Figure(figsize=(6.5, 3.75), dpi=70)
        subplot = figure.add_subplot(111)

        # ดึงและรวมค่า "คาร์บอน" จากข้อมูลโปรไฟล์ 1 และ 2
        x_labels = [profile1, profile2]
        y_values = [0, 0]

        # รวมค่าคาร์บอนสำหรับโปรไฟล์ 1
        for item in self.profile1_treeview.get_children():
            y_values[0] += float(self.profile1_treeview.item(item, "values")[1])  # สมมติว่า "Carbon" เป็นคอลัมน์ที่สอง

        # รวมค่าคาร์บอนสำหรับโปรไฟล์ 2
        for item in self.profile2_treeview.get_children():
            y_values[1] += float(self.profile2_treeview.item(item, "values")[1])  # สมมติว่า "Carbon" เป็นคอลัมน์ที่สอง

        # สร้างกราฟแท่งโดยใช้สีต่าง ๆ สำหรับแต่ละแท่ง
        bars = subplot.bar(range(len(x_labels)), y_values, color=['#C0E4F6', '#E8ABB5'], width=0.1, align='center')

        # เพิ่มตัวเลขกำกับที่กราฟแท่ง
        for bar in bars:
            yval = bar.get_height()
            subplot.text(bar.get_x() + bar.get_width() / 2, yval + 0.05, f'{yval:.2f}', ha='center', va='bottom', fontsize=8, fontweight='bold', color='black')

        # เพิ่มหัวเรื่อง
        subplot.set_title('Comparison of Carbon footprint', fontsize=8, fontweight='bold', color='black')
        subplot.set_ylabel('KgCO2eq', fontsize=8, fontweight='bold', color='black')

        # ปรับแต่งเส้นขอบ
        subplot.spines['top'].set_visible(False)
        subplot.spines['right'].set_visible(False)
        subplot.spines['left'].set_color('black')
        subplot.spines['bottom'].set_color('black')

        # กำหนดแบบอักษรและสีของเส้น
        subplot.tick_params(axis='both', which='major', labelsize=8, colors='black')

         # กำหนดระยะห่างระหว่างแท่ง
        subplot.set_xticks(range(len(x_labels)))

        # กำหนดแท่งใหม่เป็นชื่อของแต่ละ profile
        subplot.set_xticklabels(x_labels)

        # สร้างเฟรมสำหรับกราฟ
        frame = tk.Frame(self, highlightbackground='black', highlightthickness=1, borderwidth=1, relief="ridge")
        frame.grid(row=0, column=0, sticky="NSWE")

        # สร้างวิดเจ็ต FigureCanvasTkAgg
        canvas = FigureCanvasTkAgg(figure, master=frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # สร้างข้อความเปรียบเทียบ
        summary_text = ""

        # เปรียบเทียบคาร์บอนฟุตพริ้นท์และสร้างข้อความคำแนะนำ
        if y_values[0] < y_values[1]:
            summary_text = f"{profile1} < {profile2}\n"
            summary_text += f"You should select {profile1} for the lowest carbon footprint."
        elif y_values[0] > y_values[1]:
            summary_text = f"{profile2} < {profile1}\n"
            summary_text += f"You should select {profile2} for the lowest carbon footprint."
        else:
            # summary_text = f"คาร์บอนฟุตพริ้นท์ของ {profile1} และ {profile2} เท่ากัน\n\n"
            summary_text += f"Carbon footprints of {profile1} and {profile2} are equal."

        # อัพเดตป้ายข้อความสรุป
        self.label_compare_summary.config(text=summary_text) # , background='#FFD10A'

    def setBreakevenGraph(self, profile1, profile2, total_cost1, total_cost2, profit1, profit2):
        # สร้างภาพ Matplotlib
        figure = Figure(figsize=(6.5, 3.75), dpi=70)
        subplot = figure.add_subplot(111)

        # ป้ายกำกับข้อมูลสำหรับแกน x
        x_labels = [profile1, profile2]

        # ค่าของแกน y สำหรับต้นทุนรวมและกำไร
        y_values = [[total_cost1, profit1], [total_cost2, profit2]]

        # กำหนดตำแหน่ง x ของแต่ละแท่งกราฟ
        x = range(len(x_labels))

        # ความกว้างของแท่งกราฟ
        bar_width = 0.25

        # สร้างกราฟแท่ง
        bars1 = subplot.bar([pos - bar_width/2 for pos in x], [y[0] for y in y_values], bar_width, label='Total Cost', color='#C0E4F6')
        bars2 = subplot.bar([pos + bar_width/2 for pos in x], [y[1] for y in y_values], bar_width, label='Profit', color='#E8ABB5')

        # เพิ่มตัวเลขกำกับที่กราฟแท่งในหน่วยเงินบาท
        for bars in [bars1, bars2]:
            for bar in bars:
                yval = bar.get_height()
                formatted_value = self.format_currency(yval)
                subplot.text(bar.get_x() + bar.get_width() / 2, yval + 0.05, formatted_value, ha='center', va='bottom', fontsize=8, fontweight='bold', color='black')

        # เพิ่มหัวเรื่องให้กับกราฟ
        subplot.set_title('Comparison of Total Cost and Profit', fontsize=10, fontweight='bold', color='black')

        # กำหนดป้ายชื่อแกน y
        subplot.set_ylabel('Baht', fontsize=8, fontweight='bold', color='black')

        # ปรับแต่งเส้นขอบของกราฟ
        subplot.spines['top'].set_visible(False)
        subplot.spines['right'].set_visible(False)
        subplot.spines['left'].set_color('black')
        subplot.spines['bottom'].set_color('black')

        # กำหนดพารามิเตอร์ของเส้นแบ่ง
        subplot.tick_params(axis='both', which='major', labelsize=8, colors='black')

        # กำหนดตำแหน่งและป้ายกำกับของ x-tick
        subplot.set_xticks(x)
        subplot.set_xticklabels(x_labels)

        # เพิ่มตำนาน (legend)
        subplot.legend()

        # สร้างเฟรมเพื่อถือกราฟ
        frame = tk.Frame(self, highlightbackground='black', highlightthickness=1, borderwidth=1, relief="ridge")
        frame.grid(row=0, column=1, sticky="NSWE")

        # สร้างวิดเจ็ต FigureCanvasTkAgg และเพิ่มลงในเฟรม
        canvas = FigureCanvasTkAgg(figure, master=frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # สร้างข้อความเปรียบเทียบสำหรับต้นทุนรวม
        total_cost_summary_text = ""
        if total_cost1 > total_cost2:
            total_cost_summary_text += f"{profile1} > {profile2}\n"
            total_cost_summary_text += f"You should select {profile2} for the lowest cost."
        elif total_cost1 < total_cost2:
            total_cost_summary_text += f"{profile2} > {profile1}\n"
            total_cost_summary_text += f"You should select {profile1} for the lowest cost."
        else:
            total_cost_summary_text += f"Total cost of {profile1} and {profile2} are equal."

        # สร้างข้อความเปรียบเทียบสำหรับกำไร
        profit_summary_text = ""
        if profit1 > profit2:
            profit_summary_text += f"{profile1} > {profile2}\n"
            profit_summary_text += f"You should select {profile1} for the highest profit."
        elif profit1 < profit2:
            profit_summary_text += f"{profile2} > {profile1}\n"
            profit_summary_text += f"You should select {profile2} for the highest profit."
        else:
            # profit_summary_text += f"กำไรของ {profile1} และ {profile2} เท่ากัน\n"
            profit_summary_text += f"Profit of {profile1} and {profile2} are equal."

        # อัพเดตป้ายข้อความสรุปด้วยพื้นหลังสี #FFD10A
        self.label_total_cost_summary.config(text=total_cost_summary_text) # , background='#FFD10A'
        self.label_profit_summary.config(text=profit_summary_text)

    def show_profile(self, profile1, profile2, rawmats_1, transpots_1,
                      performances_1, rawmats_2, transpots_2, performances_2, breakpoint_data1, breakpoint_data2):
        
        # Check if any of the data is None and initialize to empty list
        self.rawmats_1 = rawmats_1 if rawmats_1 is not None else []
        self.transpots_1 = transpots_1 if transpots_1 is not None else []
        self.performances_1 = performances_1 if performances_1 is not None else []
        self.rawmats_2 = rawmats_2 if rawmats_2 is not None else []
        self.transpots_2 = transpots_2 if transpots_2 is not None else []
        self.performances_2 = performances_2 if performances_2 is not None else []

        # แสดงข้อมูลโปรไฟล์ที่เลือกใน Label
        self.label_profile1.config(text=profile1)
        self.label_profile2.config(text=profile2)
        self.break_profile1.config(text=profile1)
        self.break_profile2.config(text=profile2)

        # Clear existing data in the treeview
        self.profile1_treeview.delete(*self.profile1_treeview.get_children())
        self.profile2_treeview.delete(*self.profile2_treeview.get_children())

        # Insert raw materials data for profile 1
        for rawmat in rawmats_1:
            self.profile1_treeview.insert("", "end", values=(rawmat[0], round(float(rawmat[1]) * float(rawmat[2]), 3), "KgCO2eq"))

        # Insert transportation data for profile 1
        for transpot in transpots_1:
            self.profile1_treeview.insert("", "end", values=(transpot[0], round(float(transpot[1]) * float(transpot[2]), 3), "KgCO2eq"))

        # Insert performances data for profile 1
        for performance in performances_1:
            self.profile1_treeview.insert("", "end", values=(performance[0], round(float(performance[1]) * float(performance[2]), 3), "KgCO2eq"))

        # Insert raw materials data for profile 2
        for rawmat in rawmats_2:
            self.profile2_treeview.insert("", "end", values=(rawmat[0], round(float(rawmat[1]) * float(rawmat[2]), 3), "KgCO2eq"))

        # Insert transportation data for profile 2
        for transpot in transpots_2:
            self.profile2_treeview.insert("", "end", values=(transpot[0], round(float(transpot[1]) * float(transpot[2]), 3), "KgCO2eq"))

        # Insert performances data for profile 2
        for performance in performances_2:
            self.profile2_treeview.insert("", "end", values=(performance[0], round(float(performance[1]) * float(performance[2]), 3), "KgCO2eq"))

        self.setCompareGraph(profile1, profile2)
        self.set_breakpoint_data(profile1, profile2, breakpoint_data1, breakpoint_data2)
        self.updateLabel()
    
    def updateLabel(self):
        # คำนวณค่าคาร์บอนทั้งหมดสำหรับโปรไฟล์ 1
        total_carbon_profile1 = sum(float(self.profile1_treeview.item(item, "values")[1]) for item in self.profile1_treeview.get_children())

        # คำนวณค่าคาร์บอนทั้งหมดสำหรับโปรไฟล์ 2
        total_carbon_profile2 = sum(float(self.profile2_treeview.item(item, "values")[1]) for item in self.profile2_treeview.get_children())

        # คำนวณเปอร์เซ็นต์ความแตกต่าง
        if total_carbon_profile1 != 0:
            percentage_difference = ((total_carbon_profile2 - total_carbon_profile1) / total_carbon_profile1) * 100
        else:
            percentage_difference = 0  # เพื่อจัดการการหารด้วยศูนย์

        # จัดรูปแบบเปอร์เซ็นต์ความแตกต่างโดยมี "+" นำหน้าหากเป็นบวก
        percentage_text = f"{percentage_difference:+.2f}%" if percentage_difference >= 0 else f"{percentage_difference:.2f}%"

        # อัปเดตป้ายข้อความด้วยเปอร์เซ็นต์ความแตกต่างและเปลี่ยนสีถ้าเป็นลบ
        if percentage_difference < 0:
            self.label_cf.config(text=percentage_text, foreground = "green")
        else:
            self.label_cf.config(text=percentage_text, foreground = "red")

    def process_item_data_profile1(self, rawmats_1, transpots_1, performances_1):
        profile_data1 = []
        for item in rawmats_1:
            name, carbon_per, amount, unit = item
            carbon_footprint = round(float(carbon_per) * float(amount), 2)
            profile_data1.append((name, carbon_per, amount, unit, carbon_footprint))

        for item in transpots_1:
            name, carbon_per, amount, unit = item
            carbon_footprint = round(float(carbon_per) * float(amount), 2)
            profile_data1.append((name, carbon_per, amount, unit, carbon_footprint))

        for item in performances_1:
            name, carbon_per, amount, unit = item
            carbon_footprint = round(float(carbon_per) * float(amount), 2)
            profile_data1.append((name, carbon_per, amount, unit, carbon_footprint))

        return profile_data1

    def process_item_data_profile2(self, rawmats_2, transpots_2, performances_2):
        profile_data2 = []
        for item in rawmats_2:
            name, carbon_per, amount, unit = item
            carbon_footprint = round(float(carbon_per) * float(amount), 2)
            profile_data2.append((name, carbon_per, amount, unit, carbon_footprint))

        for item in transpots_2:
            name, carbon_per, amount, unit = item
            carbon_footprint = round(float(carbon_per) * float(amount), 2)
            profile_data2.append((name, carbon_per, amount, unit, carbon_footprint))

        for item in performances_2:
            name, carbon_per, amount, unit = item
            carbon_footprint = round(float(carbon_per) * float(amount), 2)
            profile_data2.append((name, carbon_per, amount, unit, carbon_footprint))

        return profile_data2
    
    def add_difference_percentage_excel(self, sheet, row, total_carbon_profile1, total_carbon_profile2):
        align_center = self.create_styles()
        
        # คำนวณเปอร์เซ็นต์ความแตกต่าง
        if total_carbon_profile1 != 0:
            percentage_difference = ((total_carbon_profile2 - total_carbon_profile1) / total_carbon_profile1) * 100
        else:
            percentage_difference = 0  # เพื่อจัดการการหารด้วยศูนย์

        # จัดรูปแบบเปอร์เซ็นต์ความแตกต่างโดยมี "+" นำหน้าหากเป็นบวก
        percentage_text = f"{percentage_difference:+.2f}" if percentage_difference >= 0 else f"{percentage_difference:.2f}"

        # เพิ่มค่าความแตกต่างนี้ลงในชีท
        sheet[f"A{row}"] = "Percentage Difference of Carbon Footprint"
        sheet[f"A{row}"].font = Font(name='TH Sarabun New', size=12)
        cell_value = sheet[f"B{row}"]
        cell_value.value = percentage_text
        cell_value.font = Font(name='TH Sarabun New', size=12)
        cell_value.number_format = '0.00%'
        cell_suffix = sheet[f"C{row}"]
        cell_suffix.value = '%'
        cell_suffix.font = Font(name='TH Sarabun New', size=12)
        cell_suffix.alignment = align_center

    def add_difference_percentage_docx(self, document, total_carbon_profile1, total_carbon_profile2):
        # คำนวณเปอร์เซ็นต์ความแตกต่าง
        if total_carbon_profile1 != 0:
            percentage_difference = ((total_carbon_profile2 - total_carbon_profile1) / total_carbon_profile1) * 100
        else:
            percentage_difference = 0  # เพื่อจัดการการหารด้วยศูนย์

        # จัดรูปแบบเปอร์เซ็นต์ความแตกต่างโดยมี "+" นำหน้าหากเป็นบวก
        percentage_text = f"{percentage_difference:+.2f}" if percentage_difference >= 0 else f"{percentage_difference:.2f}"

        # เพิ่มข้อมูลความแตกต่างในเอกสาร
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        run_label = p.add_run('Percentage Difference of Carbon Footprint: ')
        run_label.font.name = 'TH Sarabun New'
        run_label.font.size = Pt(16)
        run_label.bold = True
        run_label.font.color.rgb = RGBColor(0, 0, 0)

        tab = p.add_run('\t\t\t')
        tab.font.name = 'TH Sarabun New'
        tab.font.size = Pt(16)

        run_percentage = p.add_run(percentage_text)
        run_percentage.font.name = 'TH Sarabun New'
        run_percentage.font.size = Pt(16)
        run_percentage.font.color.rgb = RGBColor(0, 0, 0)

        tab = p.add_run('\t\t\t\t\t\t\t')
        tab.font.name = 'TH Sarabun New'
        tab.font.size = Pt(16)

        run_unit = p.add_run('%')
        run_unit.font.name = 'TH Sarabun New'
        run_unit.font.size = Pt(16)
        run_unit.font.color.rgb = RGBColor(0, 0, 0)
    
    def calculate_totals(self, data):
        return round(sum(item[4] for item in data), 2)
    
    def create_styles(self):
        align_center = Alignment(horizontal="center", vertical="center")
        return align_center

    def add_financial_summary_1(self, sheet, summary_row):
        align_center = self.create_styles()
        financial_data_1 = [
            ("Total cost", self.total_cost1, '฿#,##0.00', 'Baht'),
            ("Revenue", self.revenue1, '฿#,##0.00', 'Baht'),
            ("Profit", self.profit1, '฿#,##0.00', 'Baht'),
            ("Break-even Point", self.breakeven1, '0.00', 'Unit'),
            ("Product Efficiency", self.product_efficiency1, '0.00', '%')]

        for i, (label, value, num_format, suffix) in enumerate(financial_data_1, start=1):
            row = summary_row + i + 1
            sheet[f"A{row}"] = label
            sheet[f"A{row}"].font = Font(name='TH Sarabun New', size=12)
            cell_value = sheet[f"B{row}"]
            cell_value.value = value
            cell_value.font = Font(name='TH Sarabun New', size=12)
            cell_value.number_format = num_format
            cell_suffix = sheet[f"C{row}"]
            cell_suffix.value = suffix
            cell_suffix.font = Font(name='TH Sarabun New', size=12)
            cell_suffix.alignment = align_center

    def add_financial_summary_2(self, sheet, summary_row):
        align_center = self.create_styles()
        financial_data_2 = [
            ("Total cost", self.total_cost2, '฿#,##0.00', 'Baht'),
            ("Revenue", self.revenue2, '฿#,##0.00', 'Baht'),
            ("Profit", self.profit2, '฿#,##0.00', 'Baht'),
            ("Break-even Point", self.breakeven2, '0.00', 'Unit'),
            ("Product Efficiency", self.product_efficiency2, '0.00', '%')]

        for i, (label, value, num_format, suffix) in enumerate(financial_data_2, start=1):
            row = summary_row + i + 1
            sheet[f"A{row}"] = label
            sheet[f"A{row}"].font = Font(name='TH Sarabun New', size=12)
            cell_value = sheet[f"B{row}"]
            cell_value.value = value
            cell_value.font = Font(name='TH Sarabun New', size=12)
            cell_value.number_format = num_format
            cell_suffix = sheet[f"C{row}"]
            cell_suffix.value = suffix
            cell_suffix.font = Font(name='TH Sarabun New', size=12)
            cell_suffix.alignment = align_center

    def add_diff_financial(self, sheet, summary_row):
        align_center = self.create_styles()

        def calculate_percentage_diff(value1, value2):
            if value1 != 0:
                return ((value2 - value1) / value1) * 100
            else:
                return 0

        financial_diff_data = [
            ("Total cost", self.total_cost1, self.total_cost2),
            ("Profit", self.profit1, self.profit2)]

        for i, (label, value1, value2) in enumerate(financial_diff_data, start=1):
            percentage_diff = calculate_percentage_diff(value1, value2)
            percentage_text = f"{percentage_diff:+.2f}" if percentage_diff >= 0 else f"{percentage_diff:.2f}"

            row = summary_row + i + 1
            sheet[f"A{row}"] = f"Percentage Difference of {label}"
            sheet[f"A{row}"].font = Font(name='TH Sarabun New', size=12)
            cell_value = sheet[f"B{row}"]
            cell_value.value = percentage_text
            cell_value.font = Font(name='TH Sarabun New', size=12)
            cell_value.number_format = '0.00%'
            cell_suffix = sheet[f"C{row}"]
            cell_suffix.value = '%'
            cell_suffix.font = Font(name='TH Sarabun New', size=12)
            cell_suffix.alignment = align_center

    def add_diff_financial_docx(self, document, total_cost1, total_cost2, profit1, profit2):
        def calculate_percentage_diff(value1, value2):
            if value1 != 0:
                return ((value2 - value1) / value1) * 100
            else:
                return 0

        financial_diff_data = [
            ("Total cost", total_cost1, total_cost2),
            ("Profit", profit1, profit2)]

        for label, value1, value2 in financial_diff_data:
            percentage_diff = calculate_percentage_diff(value1, value2)
            percentage_text = f"{percentage_diff:+.2f}%" if percentage_diff >= 0 else f"{percentage_diff:.2f}%"

            # Add summary data in a separate paragraph for each label
            p = document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            run_label = p.add_run(f"Percentage Difference of {label}: ")
            run_label.font.name = 'TH Sarabun New'
            run_label.font.size = Pt(16)
            run_label.bold = True
            run_label.font.color.rgb = RGBColor(0, 0, 0)

            # Adjusting the tabs to ensure consistent spacing
            tab = p.add_run('\t\t\t\t')
            tab.font.name = 'TH Sarabun New'
            tab.font.size = Pt(16)

            run_value = p.add_run(percentage_text)
            run_value.font.name = 'TH Sarabun New'
            run_value.font.size = Pt(16)
            run_value.font.color.rgb = RGBColor(0, 0, 0)

            tab = p.add_run('\t\t\t\t\t\t')
            tab.font.name = 'TH Sarabun New'
            tab.font.size = Pt(16)

            run_unit = p.add_run('%')
            run_unit.font.name = 'TH Sarabun New'
            run_unit.font.size = Pt(16)
            run_unit.font.color.rgb = RGBColor(0, 0, 0)

    def set_column_widths(self, sheet):
        for column_cells in sheet.columns:
            max_length = max(len(str(cell.value)) for cell in column_cells if cell.value)
            adjusted_width = max_length + 2
            sheet.column_dimensions[get_column_letter(column_cells[0].column)].width = adjusted_width

    def set_page_layout(self, sheet):
        sheet.page_setup.fitToWidth = 1
        sheet.page_setup.fitToHeight = 0
        sheet.page_setup.paperSize = sheet.PAPERSIZE_A4

    def export(self):
        if any(attr is None for attr in [self.total_cost1, self.revenue1, self.profit1, self.breakeven1, self.product_efficiency1,
                                        self.total_cost2, self.revenue2, self.profit2, self.breakeven2, self.product_efficiency2]):
            print("No data available for export")
            return

        # Get profile names from label_profile1 and label_profile2
        profile1 = self.label_profile1.cget("text")
        profile2 = self.label_profile2.cget("text")

        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("Word files", "*.docx")],
            title="Save file as")

        if file_path:
            if file_path.endswith('.xlsx'):
                self.export_excel(file_path, profile1, profile2)
            elif file_path.endswith('.docx'):
                self.export_docx(file_path, profile1, profile2)

    def export_excel(self, file_path, profile1, profile2):
        profile_data1 = self.process_item_data_profile1(self.rawmats_1, self.transpots_1, self.performances_1)
        profile_data2 = self.process_item_data_profile2(self.rawmats_2, self.transpots_2, self.performances_2)
        
        total_input_carbon_footprint_1 = self.calculate_totals(profile_data1)
        total_input_carbon_footprint_2 = self.calculate_totals(profile_data2)
        
        align_center = self.create_styles()

        wb = Workbook()
        sheet = wb.active
        sheet.title = "Comparison Report"

        # Create report header
        sheet.merge_cells("A1:E1")
        sheet["A1"].value = "Comparison Report"
        sheet["A1"].alignment = align_center
        sheet["A1"].font = Font(name='TH Sarabun New', size=14, bold=True)

        # Create project name header for profile1
        sheet.merge_cells("A2:E2")
        sheet["A2"].value = f"Project Name: {profile1}"
        sheet["A2"].font = Font(name='TH Sarabun New', size=12, bold=True)

        # Create table headers
        headers = ["Name", "Emission Factor", "Amount", "Unit", "Carbon Footprint"]
        for col_num, header in enumerate(headers, start=1):
            cell = sheet.cell(row=3, column=col_num)
            cell.value = header
            cell.alignment = align_center
            cell.font = Font(name='TH Sarabun New', size=12, bold=True)

        # Fill in data rows for profile 1
        row_num = 4
        for data in profile_data1:
            for col_num, value in enumerate(data, start=1):
                cell = sheet.cell(row=row_num, column=col_num)
                cell.value = value
                cell.font = Font(name='TH Sarabun New', size=12)
                if col_num in [2, 3, 5]:
                    cell.number_format = '0.00'
            row_num += 1

        # Add summary data for profile 1
        summary_row_1 = row_num + 1
        sheet[f"A{summary_row_1}"] = "Total Carbon Footprint"
        sheet[f"A{summary_row_1}"].font = Font(name='TH Sarabun New', size=12)
        sheet[f"B{summary_row_1}"] = total_input_carbon_footprint_1
        sheet[f"B{summary_row_1}"].font = Font(name='TH Sarabun New', size=12)
        sheet[f"B{summary_row_1}"].number_format = '0.00'
        sheet[f"C{summary_row_1}"] = "KgCO2eq"
        sheet[f"C{summary_row_1}"].alignment = align_center
        sheet[f"C{summary_row_1}"].font = Font(name='TH Sarabun New', size=12)

        self.add_financial_summary_1(sheet, summary_row_1)

        # Create project name header for profile2
        row_num = summary_row_1 + 8
        sheet.merge_cells(f"A{row_num}:E{row_num}")
        sheet[f"A{row_num}"].value = f"Project Name: {profile2}"
        sheet[f"A{row_num}"].font = Font(name='TH Sarabun New', size=12, bold=True)

        # Create table headers for profile2
        row_num += 1
        for col_num, header in enumerate(headers, start=1):
            cell = sheet.cell(row=row_num, column=col_num)
            cell.value = header
            cell.alignment = align_center
            cell.font = Font(name='TH Sarabun New', size=12, bold=True)

        # Fill in data rows for profile 2
        row_num += 1
        for data in profile_data2:
            for col_num, value in enumerate(data, start=1):
                cell = sheet.cell(row=row_num, column=col_num)
                cell.value = value
                cell.font = Font(name='TH Sarabun New', size=12)
                if col_num in [2, 3, 5]:
                    cell.number_format = '0.00'
            row_num += 1

        # Add summary data for profile 2
        summary_row_2 = row_num + 1
        sheet[f"A{summary_row_2}"] = "Total Carbon Footprint"
        sheet[f"A{summary_row_2}"].font = Font(name='TH Sarabun New', size=12)
        sheet[f"B{summary_row_2}"] = total_input_carbon_footprint_2
        sheet[f"B{summary_row_2}"].font = Font(name='TH Sarabun New', size=12)
        sheet[f"B{summary_row_2}"].number_format = '0.00'
        sheet[f"C{summary_row_2}"] = "KgCO2eq"
        sheet[f"C{summary_row_2}"].alignment = align_center
        sheet[f"C{summary_row_2}"].font = Font(name='TH Sarabun New', size=12)

        self.add_financial_summary_2(sheet, summary_row_2)

         # Add percentage difference for total carbon footprint
        self.add_difference_percentage_excel(sheet, summary_row_2 + 7, total_input_carbon_footprint_1, total_input_carbon_footprint_2)

        # Add percentage difference for financial data
        self.add_diff_financial(sheet, summary_row_2 + 6)

        self.set_column_widths(sheet)
        self.set_page_layout(sheet)

        # Save the workbook
        wb.save(file_path)
        print(f"File saved at: {file_path}")

    def export_docx(self, file_path, profile1, profile2):
        document = Document()

        # Set landscape orientation
        section = document.sections[-1]
        section.orientation = WD_ORIENTATION.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Add report heading
        heading = document.add_heading('Comparison Report', level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in heading.runs:
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True

        # Add project name heading for profile1
        subheading = document.add_heading(f"Project Name: {profile1}", level=1)
        for run in subheading.runs:
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True

        # Process data for profile1
        profile_data1 = self.process_item_data_profile1(self.rawmats_1, self.transpots_1, self.performances_1)
        total_input_carbon_footprint_1 = self.calculate_totals(profile_data1)

        # Add table for profile1 data
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        headers = ["Name", "Emission Factor", "Amount", "Unit", "Carbon Footprint"]
        for cell, header in zip(hdr_cells, headers):
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'TH Sarabun New'
                    run.font.size = Pt(18)
                    run.bold = True

        for item in profile_data1:
            row_cells = table.add_row().cells
            for cell, value in zip(row_cells, item):
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'TH Sarabun New'
                        run.font.size = Pt(16)

        # Add summary for profile1
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        run_left = p.add_run('Total Carbon Footprint:')
        run_left.font.name = 'TH Sarabun New'
        run_left.font.size = Pt(16)
        run_left.bold = True
        run_left.font.color.rgb = RGBColor(0, 0, 0)
        
        tab = p.add_run('\t\t\t\t')
        tab.font.name = 'TH Sarabun New'
        tab.font.size = Pt(16)
        
        run_center = p.add_run(f'{total_input_carbon_footprint_1}')
        run_center.font.name = 'TH Sarabun New'
        run_center.font.size = Pt(16)
        run_center.font.color.rgb = RGBColor(0, 0, 0)
        
        tab = p.add_run('\t\t\t\t\t\t')
        tab.font.name = 'TH Sarabun New'
        tab.font.size = Pt(16)
        
        run_right = p.add_run('KgCO2eq')
        run_right.font.name = 'TH Sarabun New'
        run_right.font.size = Pt(16)
        run_right.font.color.rgb = RGBColor(0, 0, 0)

        # Add financial summary for profile1
        document.add_page_break()
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run('Financial Summary')
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(18)
        run.bold = True

        # Add financial summary in a table for profile1
        financial_data_1 = [
            ("Total cost", f"{self.total_cost1:,.2f}", 'Baht'),
            ("Revenue", f"{self.revenue1:,.2f}", 'Baht'),
            ("Profit", f"{self.profit1:,.2f}", 'Baht'),
            ("Break-even Point", f"{self.breakeven1:.2f}", 'Units'),
            ("Product Efficiency", f"{self.product_efficiency1:.2f}", '%')]

        table = document.add_table(rows=0, cols=3)
        for label, value, suffix in financial_data_1:
            row_cells = table.add_row().cells
            label_cell = row_cells[0].paragraphs[0].add_run(label)
            label_cell.font.name = 'TH Sarabun New'
            label_cell.font.size = Pt(16)
            label_cell.bold = True

            value_paragraph = row_cells[1].paragraphs[0]
            value_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            value_cell = value_paragraph.add_run(value)
            value_cell.font.name = 'TH Sarabun New'
            value_cell.font.size = Pt(16)

            suffix_paragraph = row_cells[2].paragraphs[0]
            suffix_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            suffix_cell = suffix_paragraph.add_run(suffix)
            suffix_cell.font.name = 'TH Sarabun New'
            suffix_cell.font.size = Pt(16)

        # Repeat the above steps for profile2
        document.add_page_break()

        # Add project name heading for profile2
        subheading = document.add_heading(f"Project Name: {profile2}", level=1)
        for run in subheading.runs:
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True

        # Process data for profile2
        profile_data2 = self.process_item_data_profile2(self.rawmats_2, self.transpots_2, self.performances_2)
        total_input_carbon_footprint_2 = self.calculate_totals(profile_data2)

        # Add table for profile2 data
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        headers = ["Name", "Emission Factor", "Amount", "Unit", "Carbon Footprint"]
        for cell, header in zip(hdr_cells, headers):
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'TH Sarabun New'
                    run.font.size = Pt(18)
                    run.bold = True

        for item in profile_data2:
            row_cells = table.add_row().cells
            for cell, value in zip(row_cells, item):
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'TH Sarabun New'
                        run.font.size = Pt(16)
        
        # Add summary for profile2
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        run_left = p.add_run('Total Carbon Footprint:')
        run_left.font.name = 'TH Sarabun New'
        run_left.font.size = Pt(16)
        run_left.bold = True
        run_left.font.color.rgb = RGBColor(0, 0, 0)
        
        tab = p.add_run('\t\t\t\t')
        tab.font.name = 'TH Sarabun New'
        tab.font.size = Pt(16)
        
        run_center = p.add_run(f'{total_input_carbon_footprint_2}')
        run_center.font.name = 'TH Sarabun New'
        run_center.font.size = Pt(16)
        run_center.font.color.rgb = RGBColor(0, 0, 0)
        
        tab = p.add_run('\t\t\t\t\t\t')
        tab.font.name = 'TH Sarabun New'
        tab.font.size = Pt(16)
        
        run_right = p.add_run('KgCO2eq')
        run_right.font.name = 'TH Sarabun New'
        run_right.font.size = Pt(16)
        run_right.font.color.rgb = RGBColor(0, 0, 0)

        # Add financial summary for profile2
        document.add_page_break()
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run('Financial Summary')
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(18)
        run.bold = True

        # Add financial summary in a table for profile2
        financial_data_2 = [
            ("Total cost", f"{self.total_cost2:,.2f}", 'Baht'),
            ("Revenue", f"{self.revenue2:,.2f}", 'Baht'),
            ("Profit", f"{self.profit2:,.2f}", 'Baht'),
            ("Break-even Point", f"{self.breakeven2:.2f}", 'Units'),
            ("Product Efficiency", f"{self.product_efficiency2:.2f}", '%')]

        table = document.add_table(rows=0, cols=3)
        for label, value, suffix in financial_data_2:
            row_cells = table.add_row().cells
            label_cell = row_cells[0].paragraphs[0].add_run(label)
            label_cell.font.name = 'TH Sarabun New'
            label_cell.font.size = Pt(16)
            label_cell.bold = True

            value_paragraph = row_cells[1].paragraphs[0]
            value_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            value_cell = value_paragraph.add_run(value)
            value_cell.font.name = 'TH Sarabun New'
            value_cell.font.size = Pt(16)

            suffix_paragraph = row_cells[2].paragraphs[0]
            suffix_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            suffix_cell = suffix_paragraph.add_run(suffix)
            suffix_cell.font.name = 'TH Sarabun New'
            suffix_cell.font.size = Pt(16)

        # เพิ่มความแตกต่างในเปอร์เซ็นต์ของคาร์บอนฟุตพริ้นท์ทั้งหมด
        document.add_paragraph()  # เว้น 1 บรรทัด
        self.add_difference_percentage_docx(document, total_input_carbon_footprint_1, total_input_carbon_footprint_2)

        # เพิ่มความแตกต่างในเปอร์เซ็นต์ของข้อมูลทางการเงิน
        self.add_diff_financial_docx(document, self.total_cost1, self.total_cost2, self.profit1, self.profit2)


        # Save the document
        if file_path:
            document.save(file_path)
            print(f"File saved at: {file_path}")
        else:
            print("File save canceled")